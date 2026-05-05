"""
doc_parser.py — Extraction de texte brut depuis tous types de documents.

Formats supportés :
  • PDF          → pdfplumber  (texte + tables)
  • Excel / ODS  → openpyxl    (feuilles + cellules)
  • CSV          → csv stdlib  (lignes + colonnes)
  • Word (.docx) → python-docx (paragraphes + tables)
  • Texte brut   → lecture directe (utf-8, latin-1, chardet)
  • JSON         → json stdlib
  • Images       → Pillow + note (OCR non embarqué)
"""
from __future__ import annotations

import csv
import io
import json
import logging
from pathlib import Path
from typing import Any

log = logging.getLogger(__name__)

# ── Résultat normalisé ────────────────────────────────────────────────────────
ParseResult = dict[str, Any]
# {
#   "text":    str,          # texte brut complet
#   "tables":  list[list],   # [[row], ...]  (optionnel, vide si non tabulaire)
#   "pages":   int,          # nombre de pages / feuilles
#   "method":  str,          # "pdf" | "excel" | "csv" | "docx" | "text" | "json" | "image" | "binary"
#   "error":   str | None,   # message d'erreur si extraction partielle
# }

MAX_TEXT = 120_000   # ~30 pages A4 — suffisant pour le LLM + contexte Limule
MAX_CELLS = 5_000    # limite de cellules pour Excel/CSV

# ── Helpers ───────────────────────────────────────────────────────────────────

def _cap(text: str) -> str:
    return text[:MAX_TEXT] if len(text) > MAX_TEXT else text


def _table_to_text(tables: list[list[Any]]) -> str:
    lines = []
    for row in tables:
        lines.append(" | ".join(str(c).strip() for c in row if c is not None))
    return "\n".join(lines)


# ── PDF ───────────────────────────────────────────────────────────────────────

def _parse_pdf(content: bytes) -> ParseResult:
    try:
        import pdfplumber  # type: ignore
    except ImportError:
        return {"text": "", "tables": [], "pages": 0, "method": "pdf", "error": "pdfplumber non installé"}

    texts: list[str] = []
    all_tables: list[list] = []
    pages = 0
    error = None
    try:
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            pages = len(pdf.pages)
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                texts.append(page_text)
                for table in page.extract_tables() or []:
                    if table:
                        all_tables.extend(table)
                        if len(all_tables) >= MAX_CELLS:
                            break
    except Exception as exc:
        error = str(exc)
        log.warning("PDF parse error: %s", exc)

    full_text = "\n\n".join(texts)
    if all_tables:
        full_text += "\n\n=== TABLEAUX ===\n" + _table_to_text(all_tables[:MAX_CELLS])

    return {
        "text": _cap(full_text.strip()),
        "tables": all_tables[:MAX_CELLS],
        "pages": pages,
        "method": "pdf",
        "error": error,
    }


# ── Excel / ODS ───────────────────────────────────────────────────────────────

def _parse_excel(content: bytes, filename: str = "") -> ParseResult:
    try:
        import openpyxl  # type: ignore
    except ImportError:
        return {"text": "", "tables": [], "pages": 0, "method": "excel", "error": "openpyxl non installé"}

    texts: list[str] = []
    all_rows: list[list] = []
    sheets = 0
    error = None

    try:
        wb = openpyxl.load_workbook(io.BytesIO(content), data_only=True, read_only=True)
        sheets = len(wb.sheetnames)
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            sheet_rows: list[list] = []
            texts.append(f"=== Feuille : {sheet_name} ===")
            for row in ws.iter_rows(values_only=True):
                clean_row = [c for c in row if c is not None]
                if clean_row:
                    sheet_rows.append(list(row))
                    texts.append(" | ".join(str(v).strip() for v in row if v is not None))
                    if len(all_rows) >= MAX_CELLS:
                        break
            all_rows.extend(sheet_rows)
    except Exception as exc:
        error = str(exc)
        log.warning("Excel parse error (%s): %s", filename, exc)

    return {
        "text": _cap("\n".join(texts).strip()),
        "tables": all_rows[:MAX_CELLS],
        "pages": sheets,
        "method": "excel",
        "error": error,
    }


# ── CSV ───────────────────────────────────────────────────────────────────────

def _parse_csv(content: bytes) -> ParseResult:
    try:
        import chardet  # type: ignore
        detected = chardet.detect(content)
        encoding = detected.get("encoding") or "utf-8"
    except ImportError:
        encoding = "utf-8"

    text_raw = content.decode(encoding, errors="replace")
    lines: list[list] = []
    try:
        dialect = csv.Sniffer().sniff(text_raw[:4096], delimiters=";,\t|")
        reader = csv.reader(io.StringIO(text_raw), dialect)
    except csv.Error:
        reader = csv.reader(io.StringIO(text_raw))

    rows_text: list[str] = []
    for row in reader:
        lines.append(row)
        rows_text.append(" | ".join(row))
        if len(lines) >= MAX_CELLS:
            break

    return {
        "text": _cap("\n".join(rows_text)),
        "tables": lines,
        "pages": 1,
        "method": "csv",
        "error": None,
    }


# ── Word (.docx) ──────────────────────────────────────────────────────────────

def _parse_docx(content: bytes) -> ParseResult:
    try:
        import docx  # type: ignore  (python-docx)
    except ImportError:
        return {"text": "", "tables": [], "pages": 0, "method": "docx", "error": "python-docx non installé"}

    texts: list[str] = []
    all_table_rows: list[list] = []
    error = None
    try:
        doc = docx.Document(io.BytesIO(content))
        for para in doc.paragraphs:
            if para.text.strip():
                texts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                all_table_rows.append(cells)
                texts.append(" | ".join(cells))
    except Exception as exc:
        error = str(exc)
        log.warning("DOCX parse error: %s", exc)

    return {
        "text": _cap("\n".join(texts).strip()),
        "tables": all_table_rows,
        "pages": len(doc.paragraphs) // 30 + 1 if not error else 0,
        "method": "docx",
        "error": error,
    }


# ── Texte brut ────────────────────────────────────────────────────────────────

def _parse_text(content: bytes) -> ParseResult:
    try:
        import chardet  # type: ignore
        detected = chardet.detect(content[:10_000])
        encoding = detected.get("encoding") or "utf-8"
    except ImportError:
        encoding = "utf-8"

    text = content.decode(encoding, errors="replace")
    return {
        "text": _cap(text),
        "tables": [],
        "pages": max(1, len(text) // 3000),
        "method": "text",
        "error": None,
    }


# ── JSON ──────────────────────────────────────────────────────────────────────

def _parse_json(content: bytes) -> ParseResult:
    try:
        data = json.loads(content.decode("utf-8", errors="replace"))
        text = json.dumps(data, ensure_ascii=False, indent=2)
    except Exception as exc:
        text = content.decode("utf-8", errors="replace")
        return {"text": _cap(text), "tables": [], "pages": 1, "method": "json", "error": str(exc)}
    return {"text": _cap(text), "tables": [], "pages": 1, "method": "json", "error": None}


# ── Image ─────────────────────────────────────────────────────────────────────

def _parse_image(content: bytes, mime_type: str) -> ParseResult:
    """
    Tente un OCR basique via pytesseract si disponible.
    Sinon retourne une note indiquant qu'une image a été reçue.
    """
    try:
        import pytesseract  # type: ignore
        from PIL import Image  # type: ignore
        img = Image.open(io.BytesIO(content))
        text = pytesseract.image_to_string(img, lang="fra+eng")
        return {"text": _cap(text.strip()), "tables": [], "pages": 1, "method": "ocr", "error": None}
    except ImportError:
        pass
    except Exception as exc:
        log.warning("OCR error: %s", exc)

    return {
        "text": f"[Image {mime_type} — OCR non disponible. Taille: {len(content)} octets]",
        "tables": [],
        "pages": 1,
        "method": "image",
        "error": "pytesseract non installé — OCR indisponible",
    }


# ── Point d'entrée principal ──────────────────────────────────────────────────

def extract_text(path: Path, mime_type: str = "", filename: str = "") -> ParseResult:
    """
    Extrait le texte d'un fichier quelle que soit son extension / MIME.
    Lit le fichier depuis `path` (doit exister).

    Returns ParseResult dict.
    """
    if not path.exists():
        return {"text": "", "tables": [], "pages": 0, "method": "missing", "error": f"Fichier introuvable: {path}"}

    content = path.read_bytes()
    ext = Path(filename or path.name).suffix.lower()
    mt = (mime_type or "").lower()

    # ── PDF ──
    if "pdf" in mt or ext == ".pdf":
        return _parse_pdf(content)

    # ── Excel ──
    if ext in {".xlsx", ".xls", ".xlsm", ".ods"} or "spreadsheet" in mt or "excel" in mt or "opendocument.spreadsheet" in mt:
        return _parse_excel(content, filename)

    # ── CSV ──
    if ext == ".csv" or "text/csv" in mt:
        return _parse_csv(content)

    # ── Word ──
    if ext in {".docx", ".doc"} or "wordprocessing" in mt or "msword" in mt:
        return _parse_docx(content)

    # ── JSON ──
    if ext == ".json" or "application/json" in mt:
        return _parse_json(content)

    # ── Images ──
    if mt.startswith("image/") or ext in {".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tiff", ".webp"}:
        return _parse_image(content, mime_type)

    # ── Texte brut / fallback ──
    if mt.startswith("text/") or ext in {".txt", ".md", ".log", ".xml", ".html", ".htm"}:
        return _parse_text(content)

    # Tentative de décodage UTF-8 pour tout le reste
    try:
        text = content.decode("utf-8", errors="strict")[:MAX_TEXT]
        return {"text": text, "tables": [], "pages": 1, "method": "text_fallback", "error": None}
    except UnicodeDecodeError:
        pass

    return {
        "text": f"[Fichier binaire — {ext or mime_type} — {len(content)} octets]",
        "tables": [],
        "pages": 0,
        "method": "binary",
        "error": "Format binaire non lisible",
    }


def extract_text_from_bytes(
    content: bytes,
    mime_type: str = "",
    filename: str = "",
) -> ParseResult:
    """
    Variante qui reçoit directement les bytes (pas besoin d'un fichier sur disque).
    Utilisée lors de l'upload avant persistance.
    """
    import tempfile

    ext = Path(filename).suffix.lower() if filename else ""
    suffix = ext or ".bin"

    with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
        tmp.write(content)
        tmp_path = Path(tmp.name)

    try:
        result = extract_text(tmp_path, mime_type=mime_type, filename=filename)
    finally:
        tmp_path.unlink(missing_ok=True)

    return result
