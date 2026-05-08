"""
routes_legislation.py — Base législative IA.

Endpoints :
  GET    /legislation/documents               — liste des documents
  POST   /legislation/documents               — uploader un document (PDF, Excel, Word, CSV…)
  DELETE /legislation/documents/{id}          — supprimer
  POST   /legislation/documents/{id}/analyze  — analyser avec Limule
  GET    /legislation/context                 — contexte législatif assemblé pour Limule
"""

from __future__ import annotations

import json
import os
from datetime import datetime, timezone
from pathlib import Path

from fastapi import APIRouter, Depends, File, Form, HTTPException, UploadFile
from sqlalchemy import select
from sqlalchemy.orm import Session

from app.api.deps import get_current_user
from app.db.session import get_db
from app.models.domain import LegislationDocument
from app.schemas.domain import LegislationDocumentCreate, LegislationDocumentRead

router = APIRouter(tags=["legislation"])

LEGISLATION_STORAGE = Path("storage/legislation")
LEGISLATION_STORAGE.mkdir(parents=True, exist_ok=True)

ALLOWED_MIME = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/msword",
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "application/vnd.ms-excel",
    "text/csv",
    "text/plain",
}

DOC_CATEGORY_LABELS = {
    "fiscal":    "Fiscalité & impôts",
    "social":    "Droit social & CNPS",
    "commerce":  "Droit commercial",
    "finance":   "Finances publiques & banque",
    "general":   "Général",
}


def _extract_text(file_path: Path, mime_type: str) -> str:
    """Extrait le texte brut d'un document uploadé."""
    try:
        if mime_type == "application/pdf":
            try:
                import pdfplumber
                with pdfplumber.open(file_path) as pdf:
                    return "\n".join(p.extract_text() or "" for p in pdf.pages[:30])
            except Exception:
                pass
        if "spreadsheet" in mime_type or "excel" in mime_type or mime_type == "text/csv":
            try:
                import pandas as pd
                df = pd.read_excel(file_path) if "excel" in mime_type or "spreadsheet" in mime_type else pd.read_csv(file_path)
                return df.to_string(index=False)[:8000]
            except Exception:
                pass
        if "word" in mime_type or "wordprocessing" in mime_type:
            try:
                from docx import Document as DocxDocument
                doc = DocxDocument(file_path)
                return "\n".join(p.text for p in doc.paragraphs)[:8000]
            except Exception:
                pass
        # Fallback: lire comme texte brut
        return file_path.read_text(errors="replace")[:8000]
    except Exception:
        return ""


# ═══════════════════════════════════════════════════════════════════
# LIST
# ═══════════════════════════════════════════════════════════════════

@router.get("/legislation/documents", response_model=list[LegislationDocumentRead])
def list_legislation_docs(
    category: str | None = None,
    db: Session = Depends(get_db),
    current_user=Depends(get_current_user),
) -> list[LegislationDocument]:
    stmt = select(LegislationDocument).where(
        LegislationDocument.company_id == current_user.company_id
    ).order_by(LegislationDocument.created_at.desc())
    if category:
        stmt = stmt.where(LegislationDocument.doc_category == category)
    return db.scalars(stmt).all()


# ═══════════════════════════════════════════════════════════════════
# UPLOAD
# ═══════════════════════════════════════════════════════════════════

@router.post("/legislation/documents", response_model=LegislationDocumentRead, status_code=201)
async def upload_legislation_doc(
    file: UploadFile = File(...),
    title: str = Form(...),
    description: str = Form(""),
    doc_category: str = Form("general"),
    country_scope: str = Form("Congo"),
    db: Session = Depends(get_db),
    current_user=Depends(get_current_user),
) -> LegislationDocument:
    if file.content_type not in ALLOWED_MIME:
        raise HTTPException(400, f"Type de fichier non supporté : {file.content_type}. Acceptés : PDF, Word, Excel, CSV, TXT.")

    content = await file.read()
    size_bytes = len(content)
    if size_bytes > 30 * 1024 * 1024:  # 30 Mo max
        raise HTTPException(400, "Fichier trop volumineux (max 30 Mo)")

    # Sauvegarde
    safe_name = f"{int(datetime.now().timestamp())}_{file.filename}"
    dest = LEGISLATION_STORAGE / safe_name
    dest.write_bytes(content)

    # Extraction texte
    raw_text = _extract_text(dest, file.content_type or "")

    doc = LegislationDocument(
        company_id=current_user.company_id,
        title=title,
        description=description,
        filename=file.filename or safe_name,
        storage_path=str(dest),
        mime_type=file.content_type or "",
        size_bytes=size_bytes,
        doc_category=doc_category,
        country_scope=country_scope,
        raw_text=raw_text[:20000],
        analyzed=False,
        uploaded_by_user_id=current_user.id,
    )
    db.add(doc)
    db.commit()
    db.refresh(doc)
    return doc


# ═══════════════════════════════════════════════════════════════════
# DELETE
# ═══════════════════════════════════════════════════════════════════

@router.delete("/legislation/documents/{doc_id}", status_code=204)
def delete_legislation_doc(
    doc_id: int,
    db: Session = Depends(get_db),
    current_user=Depends(get_current_user),
) -> None:
    doc = db.get(LegislationDocument, doc_id)
    if not doc or doc.company_id != current_user.company_id:
        raise HTTPException(404, "Document introuvable")
    # Supprimer le fichier
    try:
        Path(doc.storage_path).unlink(missing_ok=True)
    except Exception:
        pass
    db.delete(doc)
    db.commit()


# ═══════════════════════════════════════════════════════════════════
# ANALYZE — Limule extrait les infos législatives clés
# ═══════════════════════════════════════════════════════════════════

@router.post("/legislation/documents/{doc_id}/analyze", response_model=LegislationDocumentRead)
async def analyze_legislation_doc(
    doc_id: int,
    db: Session = Depends(get_db),
    current_user=Depends(get_current_user),
) -> LegislationDocument:
    from app.services.limule import limule_generate

    doc = db.get(LegislationDocument, doc_id)
    if not doc or doc.company_id != current_user.company_id:
        raise HTTPException(404, "Document introuvable")
    if not doc.raw_text:
        raise HTTPException(400, "Aucun texte extrait — réuploadez le document.")

    cat_label = DOC_CATEGORY_LABELS.get(doc.doc_category, doc.doc_category)
    prompt = (
        f"Analyse ce document législatif de catégorie '{cat_label}' "
        f"(portée géographique : {doc.country_scope}).\n\n"
        f"TEXTE DU DOCUMENT :\n{doc.raw_text[:6000]}\n\n"
        f"Fournis :\n"
        f"1. RÉSUMÉ : synthèse des dispositions clés (obligations, seuils, taux, délais)\n"
        f"2. IMPACTS PME : ce que chaque PME doit retenir et appliquer concrètement\n"
        f"3. DATES CLÉS : échéances légales, délais de dépôt, périodicités\n"
        f"4. RISQUES : sanctions, pénalités et cas de non-conformité\n"
        f"5. TAGS : liste de 5-10 mots-clés séparés par des virgules\n"
        f"Format sections clairement avec les numéros ci-dessus."
    )

    content, _ = await limule_generate(
        kind="legislation",
        prompt=prompt,
        context="",
        db=db,
        company_id=current_user.company_id,
        user=current_user,
        max_tokens=3000,
        temperature=0.15,
    )

    # Extraire les tags de la section 5
    tags: list[str] = []
    if "5." in content or "TAGS" in content.upper():
        lines = content.split("\n")
        for i, line in enumerate(lines):
            if "5." in line or "TAGS" in line.upper():
                tag_line = lines[i + 1] if i + 1 < len(lines) else ""
                tags = [t.strip() for t in tag_line.split(",") if t.strip()][:10]
                break

    doc.ai_summary = content
    doc.ai_tags = json.dumps(tags, ensure_ascii=False)
    doc.analyzed = True
    doc.updated_at = datetime.now(timezone.utc)
    db.commit()
    db.refresh(doc)
    return doc


# ═══════════════════════════════════════════════════════════════════
# CONTEXT — contexte législatif assemblé pour injection dans Limule
# ═══════════════════════════════════════════════════════════════════

@router.get("/legislation/context")
def get_legislation_context(
    db: Session = Depends(get_db),
    current_user=Depends(get_current_user),
) -> dict:
    docs = db.scalars(
        select(LegislationDocument).where(
            LegislationDocument.company_id == current_user.company_id,
            LegislationDocument.analyzed == True,
        ).order_by(LegislationDocument.created_at.desc()).limit(10)
    ).all()

    if not docs:
        return {"context": "", "doc_count": 0}

    parts = []
    for doc in docs:
        cat = DOC_CATEGORY_LABELS.get(doc.doc_category, doc.doc_category)
        summary = (doc.ai_summary or "")[:800]
        parts.append(f"[{cat} — {doc.title}]\n{summary}")

    context_text = "\n\n---\n\n".join(parts)
    return {
        "context": context_text,
        "doc_count": len(docs),
        "categories": list({d.doc_category for d in docs}),
    }
